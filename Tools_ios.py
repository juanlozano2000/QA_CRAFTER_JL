import os
from appium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from appium.webdriver.common.appiumby import AppiumBy
from appium.webdriver.common.touch_action import TouchAction
import os
from PIL import Image
from docx import Document
from docx.shared import Inches

screenshot_folder = 'J:\Sistemas\CC1241\Vanesa Garcia\Pruebas consulto mis Tarjetas\Trimestre 27\Sprint 2\Regre IOS\IOS 9.18.0'

class Tools:
    def __init__(self, driver):
        self.driver = driver

    def click_xpath(self, xpath):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.XPATH, xpath))
        ).click()

    def click_ACCESSIBILITY_ID(self, accessibility):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.ACCESSIBILITY_ID, accessibility))
        ).click()

    def click_element_ID(self, elementId):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.ID, elementId))
        ).click()

    def click_className(self, className):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.CLASS_NAME, className))
        ).click()

    def click_id(self, element_id):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.ID, f"{element_id}"))
        ).click()

    def click_position(self, x, y):
        action = TouchAction(self.driver)
        action.tap(x = x, y = y).perform()

    def click_chain(self, chain):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.IOS_CLASS_CHAIN, chain))
        ).click()

    def set_id(self, element_id, value):
        element = WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.XPATH, element_id))
        )
        element.send_keys(value)
    
    def set_id_real(self, element_id, value):
        element = WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.ID, element_id))
        )
        element.send_keys(value)

    def expected_locator(self, xpath):
        WebDriverWait(self.driver, 15).until(
            EC.presence_of_element_located((AppiumBy.XPATH, xpath))
        )

    def expected_btn_disabled(self, xpath):
        def is_disabled():
            try:
                element = WebDriverWait(self.driver, 15).until(
                    EC.presence_of_element_located((AppiumBy.XPATH, xpath))
                )
                return element.get_attribute("enabled") == "false"
            except:
                return False
        WebDriverWait(self.driver, 10).until(is_disabled)

    def expected_btn_enabled(self, xpath):
        def is_disabled():
            try:
                element = WebDriverWait(self.driver, 15).until(
                    EC.presence_of_element_located((AppiumBy.XPATH, xpath))
                )
                return element.get_attribute("enabled") == "true"
            except:
                return False
        WebDriverWait(self.driver, 10).until(is_disabled)

    def validar_label(self, xpath, valor_esperado):
        elemento = WebDriverWait(self.driver, 15).until(
                EC.presence_of_element_located((AppiumBy.XPATH, xpath))
            )
        label = elemento.get_attribute("label")
        assert label == valor_esperado, f"Label esperado: '{valor_esperado}', pero se encontró: '{label}'"

    def not_expected_locator(self, xpath):
        WebDriverWait(self.driver, 15).until(
            EC.invisibility_of_element_located((AppiumBy.XPATH, xpath))
        )

    def expected_text(self, xpath, text):
        WebDriverWait(self.driver, 15).until(
            EC.text_to_be_present_in_element((AppiumBy.XPATH , xpath), text)
        )

    def not_expected_text(self, xpath, text):
        WebDriverWait(self.driver, 15).until_not(
            EC.text_to_be_present_in_element((AppiumBy.XPATH , xpath), text)
        )

    def expected_text_chain(self, chain, text):
        WebDriverWait(self.driver, 15).until(
            EC.text_to_be_present_in_element((AppiumBy.IOS_CLASS_CHAIN , chain), text)
        )

    def expected_contain_text(self, xpath, text):
        WebDriverWait(self.driver, 15).until(
            EC.text_to_be_present_in_element((AppiumBy.XPATH , xpath), text)
        )
    
    def take_screenshot(self, folder_path, screenshot_name):
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        screenshot_path = os.path.join(folder_path, f"{screenshot_name}.png")
        self.driver.get_screenshot_as_file(screenshot_path)
        print(f"Screenshot saved to {screenshot_path}")

    def scroll(self, start_x, start_y, end_x, end_y):
        action = TouchAction(self.driver)
        action.press(x=start_x, y=start_y).wait(ms=500).move_to(x=end_x, y=end_y).release().perform()
        print(f"Scrolled from ({start_x}, {start_y}) to ({end_x}, {end_y})")

    def getText(self, xpath):
        element = WebDriverWait(self.driver, 15).until(
            EC.visibility_of_element_located((AppiumBy.XPATH, xpath))
        )
        return element.text
    
    def achicarImgs(self, image_name, scale_factor=4.5):
        """
        Redimensiona una imagen existente en la carpeta especificada y guarda la versión reducida con el sufijo '_r'.
        """
        # Construir las rutas de entrada y salida
        img_path = os.path.join(screenshot_folder, f'{image_name}.png')
        print(f'{img_path} ruta de la imagen SKERE')
        #resized_path = os.path.join(screenshot_folder, f"{os.path.splitext(image_name)[0]}_r{os.path.splitext(image_name)[1]}")
        resized_path = os.path.join(screenshot_folder, f"{image_name}_r.png")

        # Verificar que la imagen existe
        if not os.path.exists(img_path):
            print(f"La imagen '{img_path}.png' no existe.")
            return

        # Redimensionar la imagen
        try:
            with Image.open(img_path) as img:
                new_size = (int(img.width / scale_factor), int(img.height / scale_factor))
                img_resized = img.resize(new_size, Image.Resampling.LANCZOS)
                
                # Guardar la imagen redimensionada con el sufijo '_r'
                img_resized.save(resized_path)
                print(f"Imagen '{image_name}' reducida a tamaño {new_size} y guardada como '{resized_path}'")
        except Exception as e:
            print(f"Error al redimensionar la imagen '{image_name}': {e}")

    def create_word_with_selected_images(self, output_filename, selected_images, name_client, responseBackend=None):
        """
        Crea un documento de Word que incluye imágenes seleccionadas de una carpeta.

        Args:
        - screenshot_folder (str): Ruta de la carpeta donde se encuentran las imágenes.
        - output_filename (str): Nombre del archivo de Word a crear.
        - selected_images (list): Lista de nombres (sin extensión) de las imágenes a incluir.
        """
        # Crear un nuevo documento de Word
        doc = Document()

        doc.add_heading(f"Testeado con: {name_client}", level=2)
        
        # Iterar sobre todos los archivos en la carpeta
        for filename in os.listdir(screenshot_folder):
            # Obtener el nombre del archivo sin extensión
            file_name_without_extension = os.path.splitext(filename)[0]
            
            # Solo procesar si el archivo está en la lista de seleccionados
            if file_name_without_extension in selected_images and filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                try:
                    # Ruta completa de la imagen
                    img_path = os.path.join(screenshot_folder, filename)
                    
                    # Añadir la imagen al documento de Word (sin modificar el tamaño)
                    doc.add_picture(img_path)  
                    doc.add_paragraph(filename)  # Agregar el nombre de la imagen como un título o pie de foto
                    
                    print(f"Imagen '{filename}' añadida al documento.")
                except Exception as e:
                    print(f"Error añadiendo la imagen {filename}: {e}")
        
        # Agregar opcionalmente un response backend
        if responseBackend is not None:
            doc.add_paragraph(responseBackend)

        # Construir la ruta completa del archivo de salida en la misma carpeta
        output_docx = os.path.join(screenshot_folder, output_filename)
        
        # Guardar el documento en la ruta especificada
        try:
            doc.save(output_docx)
            print(f"Documento guardado como '{os.path.abspath(output_docx)}'.")
        except Exception as e:
            print(f"Error al guardar el documento: {e}")


    def get_element_position(self, xpath):
        """
        Devuelve la posición X e Y del elemento ubicado por el XPath.

        Args:
        - xpath (str): El XPath del elemento a localizar.

        Returns:
        - dict: Un diccionario con las claves 'x' y 'y' para la posición del elemento.
        """
        try:
            # Encontrar el elemento usando XPath
            element = self.driver.find_element(AppiumBy.XPATH, xpath)
            
            # Obtener la posición del elemento
            position = element.location
            print(f"Posición del elemento '{xpath}': X={position['x']}, Y={position['y']}")
            return position
        except Exception as e:
            print(f"Error al obtener la posición del elemento '{xpath}': {e}")
            return None


    def scroll_carrousel(self, direccion):
        if direccion == "Izquierda":
            self.scroll(338,225, 79, 225)
        else:
            self.scroll(79, 225, 338,225)

    def expected_text_equal(self, text_1, text_2):
        try:
            # Usa WebDriverWait con una condición personalizada
            WebDriverWait(self.driver, 15).until(
                lambda driver: text_1 in text_2
            )
            print(f"El texto '{text_1}' está presente en '{text_2}'.")
        except Exception as e:
            raise AssertionError(f"El texto esperado '{text_1}' no está presente en '{text_2}'. Error: {e}")
        
    def log_browserstack(self, mensaje, level="Debug"):
        self.driver.execute_script(
            f'browserstack_executor: {"{"}"action": "annotate", "arguments": {{"data":"{mensaje}", "level": "{level}"}}{"}"}'
        )