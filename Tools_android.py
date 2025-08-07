import os
""" from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas """
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from appium.webdriver.common.appiumby import AppiumBy
from appium.webdriver.common.touch_action import TouchAction
from selenium.webdriver.common.actions.action_builder import ActionBuilder
from selenium.webdriver.common.actions.pointer_input import PointerInput
from PIL import Image
from docx import Document
from docx.shared import Inches

screenshot_folder = 'J:\Sistemas\CC1241\Vanesa Garcia\Pruebas consulto mis Tarjetas\Trimestre 26\Sprint 5\Mensajes de error IOS\\autoVane'


class Tools:
    def __init__(self, driver):
        self.driver = driver

    def click_xpath(self, xpath):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.XPATH, xpath))
        ).click()

    def click_ACCESSIBILITY_ID(self, accessibility):        
        WebDriverWait(self.driver, 30).until(
            EC.element_to_be_clickable((AppiumBy.ACCESSIBILITY_ID, accessibility))
        ).click()

    def click_className(self, className):
        WebDriverWait(self.driver, 30).until(
            EC.element_to_be_clickable((AppiumBy.CLASS_NAME, className))
        ).click()

    def click_id(self, element_id):
        WebDriverWait(self.driver, 15).until(
            EC.element_to_be_clickable((AppiumBy.ID, element_id))
        ).click()

    """ def click_position(self, x, y):
        action = TouchAction(self.driver)
        action.tap(x = x, y = y).perform() """

    def set_id(self, element_id, value):
        element = WebDriverWait(self.driver, 30).until(
            EC.element_to_be_clickable((AppiumBy.ID, element_id))
        )
        element.send_keys(value)

    def expected_locator(self, xpath):
        WebDriverWait(self.driver, 15).until(
            EC.presence_of_element_located((AppiumBy.XPATH, xpath))
        )

    def not_expected_locator(self, xpath):
        WebDriverWait(self.driver, 15).until(
            EC.invisibility_of_element_located((AppiumBy.XPATH, xpath))
        )

    def expected_text(self, xpath, text):
        WebDriverWait(self.driver, 15).until(
            EC.text_to_be_present_in_element((AppiumBy.XPATH , xpath), text)
        )
        
    def expected_text_equal(self, text_1, text_2):
        try:
            # Usa WebDriverWait con una condición personalizada
            WebDriverWait(self.driver, 15).until(
                lambda driver: text_1 in text_2
            )
            print(f"El texto '{text_1}' está presente en '{text_2}'.")
        except Exception as e:
            raise AssertionError(f"El texto esperado '{text_1}' no está presente en '{text_2}'. Error: {e}")

    def expected_contain_text(self, xpath, text):
        WebDriverWait(self.driver, 30).until(
            EC.text_to_be_present_in_element((AppiumBy.XPATH , xpath), text)
        )
    
    def take_screenshot(self, folder_path, screenshot_name):
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        screenshot_path = os.path.join(folder_path, f"{screenshot_name}.png")
        self.driver.get_screenshot_as_file(screenshot_path)
        print(f"Screenshot saved to {screenshot_path}")

    def scroll(self, start_x, start_y, end_x, end_y):
        action = TouchAction(self.driver)
        action.press(x=start_x, y=start_y).wait(ms=1000).move_to(x=end_x, y=end_y).release().perform()
        print(f"Scrolled from ({start_x}, {start_y}) to ({end_x}, {end_y})")

    def scrollnew(self, start_x, start_y, end_x, end_y):
        self.driver.execute_script("mobile: swipe", {
            "startX": start_x,
            "startY": start_y,
            "endX": end_x,
            "endY": end_y,
            "speed": 1000  # Opcional, define la velocidad del swipe
        })
        print(f"Scrolled from ({start_x}, {start_y}) to ({end_x}, {end_y})")

    def scrollnew2(self, start_x, start_y, end_x, end_y):
        actions = ActionBuilder(self.driver)
        finger = PointerInput(PointerInput.TOUCH, "finger")

        actions.add_action(finger.create_pointer_move(0, 'viewport', start_x, start_y))
        actions.add_action(finger.create_pointer_down())
        actions.add_action(finger.create_pause(0.2))
        actions.add_action(finger.create_pointer_move(0.2, 'viewport', end_x, end_y))
        actions.add_action(finger.create_pointer_up())

        self.driver.perform(actions)
        print(f"Scrolled from ({start_x}, {start_y}) to ({end_x}, {end_y})")

        
    def getText(self, xpath):
        element = WebDriverWait(self.driver, 15).until(
            EC.visibility_of_element_located((AppiumBy.XPATH, xpath))
        )
        return element.text
    
    def achicarImgs(self, image_name, scale_factor=4.5):
        """
        Redimensiona una imagen existente en la carpeta especificada y guarda la versión reducida con el sufijo '_r'.
        """
        # Construir las rutas de entrada y salida
        img_path = os.path.join(screenshot_folder, f'{image_name}.png')
        print(f'{img_path} ruta de la imagen SKERE')
        #resized_path = os.path.join(screenshot_folder, f"{os.path.splitext(image_name)[0]}_r{os.path.splitext(image_name)[1]}")
        resized_path = os.path.join(screenshot_folder, f"{image_name}_r.png")

        # Verificar que la imagen existe
        if not os.path.exists(img_path):
            print(f"La imagen '{img_path}.png' no existe.")
            return

        # Redimensionar la imagen
        try:
            with Image.open(img_path) as img:
                new_size = (int(img.width / scale_factor), int(img.height / scale_factor))
                img_resized = img.resize(new_size, Image.Resampling.LANCZOS)
                
                # Guardar la imagen redimensionada con el sufijo '_r'
                img_resized.save(resized_path)
                print(f"Imagen '{image_name}' reducida a tamaño {new_size} y guardada como '{resized_path}'")
        except Exception as e:
            print(f"Error al redimensionar la imagen '{image_name}': {e}")

    def create_word_with_selected_images(self, output_filename, selected_images, name_client, responseBackend=None):
        """
        Crea un documento de Word que incluye imágenes seleccionadas de una carpeta.

        Args:
        - screenshot_folder (str): Ruta de la carpeta donde se encuentran las imágenes.
        - output_filename (str): Nombre del archivo de Word a crear.
        - selected_images (list): Lista de nombres (sin extensión) de las imágenes a incluir.
        """
        # Crear un nuevo documento de Word
        doc = Document()

        doc.add_heading(f"Testeado con: {name_client}", level=2)
        
        # Iterar sobre todos los archivos en la carpeta
        for filename in os.listdir(screenshot_folder):
            # Obtener el nombre del archivo sin extensión
            file_name_without_extension = os.path.splitext(filename)[0]
            
            # Solo procesar si el archivo está en la lista de seleccionados
            if file_name_without_extension in selected_images and filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                try:
                    # Ruta completa de la imagen
                    img_path = os.path.join(screenshot_folder, filename)
                    
                    # Añadir la imagen al documento de Word (sin modificar el tamaño)
                    doc.add_picture(img_path)  
                    doc.add_paragraph(filename)  # Agregar el nombre de la imagen como un título o pie de foto
                    
                    print(f"Imagen '{filename}' añadida al documento.")
                except Exception as e:
                    print(f"Error añadiendo la imagen {filename}: {e}")
        
        # Agregar opcionalmente un response backend
        if responseBackend is not None:
            doc.add_paragraph(responseBackend)

        # Construir la ruta completa del archivo de salida en la misma carpeta
        output_docx = os.path.join(screenshot_folder, output_filename)
        
        # Guardar el documento en la ruta especificada
        try:
            doc.save(output_docx)
            print(f"Documento guardado como '{os.path.abspath(output_docx)}'.")
        except Exception as e:
            print(f"Error al guardar el documento: {e}")

    def log_browserstack(self, mensaje, level="info"):
        self.driver.execute_script(
            f'browserstack_executor: {"{"}"action": "annotate", "arguments": {{"data":"{mensaje}", "level": "{level}"}}{"}"}'
        )

    def scroll_carrousel(self, direccion):
        if direccion == "Izquierda" or "izquierda":
            self.scroll(869,668, 150, 668)
        else:
            self.scroll(150, 225, 869,225)

    def scroll_vertical_al_fondo(self):
        self.scroll(480, 1724, 480, 304)

    def scroll_vertical_hacia_arriba(self):
        self.scroll(480, 304, 480, 1724)